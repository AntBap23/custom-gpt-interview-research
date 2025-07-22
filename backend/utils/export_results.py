from docx import Document
from fpdf import FPDF
from docx2pdf import convert
import json
import os

def export_interview_to_docx(data, docx_path):
    doc = Document()
    doc.add_heading("Simulated Interview Transcript", level=1)
    for item in data:
        doc.add_heading(f"Q: {item['question']}", level=2)
        doc.add_paragraph(item['answer'])
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)

def export_interview_to_pdf(data, pdf_path):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, "Simulated Interview Transcript", ln=True, align='C')
    pdf.ln(5)
    pdf.set_font("Arial", size=12)
    for item in data:
        pdf.set_font("Arial", 'B', 12)
        pdf.multi_cell(0, 10, f"Q: {item['question']}")
        pdf.set_font("Arial", '', 12)
        pdf.multi_cell(0, 10, f"A: {item['answer']}")
        pdf.ln(4)
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    pdf.output(pdf_path)

def export_both(input_json_path, output_dir="outputs", output_basename="full_interviews"):
    with open(input_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    docx_path = os.path.join(output_dir, f"{output_basename}.docx")
    pdf_path = os.path.join(output_dir, f"{output_basename}.pdf")
    export_interview_to_docx(data, docx_path)
    export_interview_to_pdf(data, pdf_path)
    if os.name == 'nt':
        try:
            convert(docx_path, pdf_path.replace(".pdf", "_windows.pdf"))
        except Exception:
            pass
    return docx_path, pdf_path 