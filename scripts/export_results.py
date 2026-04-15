import csv
import html
import json
import os

from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos


def _safe_text(value):
    return str(value or "").replace("\r\n", "\n").replace("\r", "\n")

def export_interview_to_docx(data, docx_path):
    """Export interview data to DOCX format."""
    doc = Document()
    doc.add_heading("Simulated Interview Transcript", level=1)

    for item in data:
        doc.add_heading(f"Q: {_safe_text(item.get('question', ''))}", level=2)
        doc.add_paragraph(_safe_text(item.get("answer", "")))
    
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    return docx_path

def export_interview_to_pdf(data, pdf_path):
    """Export interview data to PDF format."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Simulated Interview Transcript", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(5)
    pdf.set_font("Helvetica", size=12)

    for item in data:
        pdf.set_font("Helvetica", "B", 12)
        question = _safe_text(item.get("question", "")).encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 10, f"Q: {question}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_font("Helvetica", "", 12)
        answer = _safe_text(item.get("answer", "")).encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 10, f"A: {answer}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)
    
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    pdf.output(pdf_path)
    return pdf_path

def export_both(input_json_path, output_basename, output_dir="outputs"):
    """Export interview data to both DOCX and PDF formats."""
    with open(input_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    os.makedirs(output_dir, exist_ok=True)
    docx_path = os.path.join(output_dir, f"{output_basename}.docx")
    pdf_path = os.path.join(output_dir, f"{output_basename}.pdf")
    
    export_interview_to_docx(data, docx_path)
    export_interview_to_pdf(data, pdf_path)
    
    return docx_path, pdf_path


def export_interview_to_csv(data, csv_path):
    """Export interview data to CSV format."""
    os.makedirs(os.path.dirname(csv_path), exist_ok=True)
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["question", "answer"])
        writer.writeheader()
        for item in data:
            writer.writerow({"question": item.get("question", ""), "answer": item.get("answer", "")})
    return csv_path


def export_interview_to_txt(data, txt_path):
    """Export interview data to plain text format."""
    os.makedirs(os.path.dirname(txt_path), exist_ok=True)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("Simulated Interview Transcript\n\n")
        for item in data:
            f.write(f"Q: {item.get('question', '')}\n")
            f.write(f"A: {item.get('answer', '')}\n\n")
    return txt_path


def export_interview_to_html(data, html_path):
    """Export interview data to a simple HTML format."""
    os.makedirs(os.path.dirname(html_path), exist_ok=True)
    body = [
        "<html><head><meta charset='utf-8'><title>Simulated Interview Transcript</title></head><body>",
        "<h1>Simulated Interview Transcript</h1>",
    ]
    for item in data:
        question = html.escape(_safe_text(item.get("question", "")))
        answer = html.escape(_safe_text(item.get("answer", ""))).replace("\n", "<br>")
        body.append(f"<h2>Q: {question}</h2>")
        body.append(f"<p>{answer}</p>")
    body.append("</body></html>")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write("\n".join(body))
    return html_path


def export_format(input_json_path, output_basename, file_type, output_dir="outputs"):
    """Export interview data to a single requested format."""
    with open(input_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    os.makedirs(output_dir, exist_ok=True)
    file_type = str(file_type or "").lower()
    output_path = os.path.join(output_dir, f"{output_basename}.{file_type}")

    exporters = {
        "docx": export_interview_to_docx,
        "pdf": export_interview_to_pdf,
        "csv": export_interview_to_csv,
        "txt": export_interview_to_txt,
        "html": export_interview_to_html,
    }
    exporter = exporters.get(file_type)
    if not exporter:
        raise ValueError(f"Unsupported export format: {file_type}")

    exporter(data, output_path)
    return output_path


def export_all_formats(input_json_path, output_basename, output_dir="outputs"):
    """Export interview data to DOCX, PDF, CSV, TXT, and HTML formats."""
    return {
        "docx": export_format(input_json_path, output_basename, "docx", output_dir=output_dir),
        "pdf": export_format(input_json_path, output_basename, "pdf", output_dir=output_dir),
        "csv": export_format(input_json_path, output_basename, "csv", output_dir=output_dir),
        "txt": export_format(input_json_path, output_basename, "txt", output_dir=output_dir),
        "html": export_format(input_json_path, output_basename, "html", output_dir=output_dir),
    }
